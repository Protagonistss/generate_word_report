#!/usr/bin/python
# -*- coding: utf-8 -*-
# @Author  : protagonisths
# @File    : generate.py
# @Time    : 2019-01-20 12:13
# @Software: PyCharm
from docx import Document
import json
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from constants.word_report import TITLE, CONTENT_TABLE_LOCATION, CONTENT_PT_FONT_SIZE, SPACE_PT_SIZE, SPACE_BEFORE_SIZE, \
    LEFT_INDENT_PT_SIZE, FIRST_LINE_INDENT_INCH_SIZE, BIG_TITLE_FONT_SIZE, FOOTER_TABLE_LINE_SPACING_SIZE, \
    FOOTER_TABLE_COLUMN_WIDTH
import logging


class GenerateWordReport(object):
    """生成word正文报告"""

    def set_column_width(self, column, width):
        """设置单元格宽度"""
        column.width = width
        for cell in column.cells:
            cell.width = width

    def set_cell_height(self, table, idx):
        """设置单元格高度"""
        if idx in [2, 5]:
            row = table.rows[idx]  # or however you get the row you want
            # print(row)
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "1200")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

    def merge_cell(self, table, idx):
        """合并单元格"""
        if idx in [0, 3, 6]:
            row = table.rows[idx]
            a, b = row.cells[:]
            a.merge(b)

    def style_content_table_row_height(self, table):
        """行高"""
        for row in table.rows:
            row.height = Cm(0.5)

    def style_content_table(self, run, idx, table, col):
        """content table style"""
        run.font.name = '宋体'
        run.font.bold = True
        # paragraph_format = run.paragraph_format
        # run.line_spacing = Pt(18)
        if (idx == 0 and col == 1) or (idx == 0 and col == 0) or (idx == 3 and col == 0) or (
                idx == 6 and col == 0) or (idx == 3 and col == 1) or (idx == 6 and col == 1):
            table.cell(idx, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            table.cell(idx, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 设置line_spacing
        for paragraph in table.cell(idx, col).paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = Pt(SPACE_PT_SIZE)

    def generate_header(self, document):
        section = document.sections[0]
        header = section.first_page_header
        paragraph = header.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Cm(-1.6)
        run = paragraph.add_run()
        run.add_picture("zzz.png", width=Cm(17.62), height=Cm(2.56))
        run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def generate_content_table(self, document):
        """生成正文中的表格"""
        with open('table.json', 'r')as f:
            content = json.load(f)
            table = document.add_table(rows=7, cols=2, style="Table Grid")
            # 行高
            self.style_content_table_row_height(table)
            table.autofit = False
            # set_column_width(table.columns[0], Cm(5))
            table_content = content.get('content')
            for idx, value in enumerate(table_content):
                self.set_cell_height(table, idx)
                for col in range(2):
                    self.merge_cell(table, idx)
                    if col == 0:
                        run = table.cell(idx, col).paragraphs[0].add_run(value.get("first"))
                    else:
                        run = table.cell(idx, col).paragraphs[0].add_run(value.get("second"))
                    # 添加table样式
                    self.style_content_table(run, idx, table, col)

    def style_footer_table_row_height(self, table):
        # 行高
        for idx, row in enumerate(table.rows):
            if idx == 4:
                row.height = Cm(0.5)
            elif idx == 3:
                row.height = Cm(0.78)
            else:
                row.height = Cm(0.88)

    def style_footer_table(self, table, idx, col, run):
        # 设置line_spacing
        for paragraph in table.cell(idx, col).paragraphs:
            paragraph_format = paragraph.paragraph_format
            if idx == 4 and (col in [0, 2]):
                paragraph_format.line_spacing = Pt(SPACE_PT_SIZE)
            else:
                paragraph_format.line_spacing = FOOTER_TABLE_LINE_SPACING_SIZE

        # table.cell(idx, col).paragraphs[0].line_spacing = Pt(18)
        # run.line_spacing = Pt(18)
        run.font.size = Pt(CONTENT_PT_FONT_SIZE)

    def generate_footer_table(self, document, sum_dict):
        """生成footer table"""
        with open('footer.json', 'r')as f:
            content = json.load(f)
            table = document.add_table(rows=5, cols=3)
            table_content = content.get('footer')
            # 行高
            self.style_footer_table_row_height(table)
            for idx, value in enumerate(table_content):
                for col in range(3):
                    if col == 0:
                        if idx == 4:
                            table.cell(idx, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            table.cell(idx, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        self.set_column_width(table.columns[col], Cm(FOOTER_TABLE_COLUMN_WIDTH))
                        if idx < 4:
                            run = table.cell(idx, col).paragraphs[0].add_run(value.get("first"))
                        else:
                            run = table.cell(idx, col).paragraphs[0].add_run(sum_dict.get("city"))
                    elif col == 1:
                        self.set_column_width(table.columns[col], Cm(0.5))
                        run = table.cell(idx, col).paragraphs[0].add_run(value.get("second"))
                        table.cell(idx, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    else:
                        self.set_column_width(table.columns[col], Cm(FOOTER_TABLE_COLUMN_WIDTH))
                        run = table.cell(idx, col).paragraphs[0].add_run(value.get("third"))
                        table.cell(idx, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    # 设置line_spacing
                    self.style_footer_table(table, idx, col, run)

    def generate_big_title(self, document, title):
        """报告开头的大标题"""
        big_title = document.add_paragraph()
        big_title_format = big_title.paragraph_format
        big_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        big_title_format.space_before = Pt(SPACE_BEFORE_SIZE)
        big_title_format.line_spacing_rule = 0
        run = big_title.add_run(title)
        run.font.bold = True
        run.font.size = Pt(BIG_TITLE_FONT_SIZE)
        run.font.name = u'SimHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'SimHei')

    def generate_head_content(self, document, content):
        """big_title具有的文字"""
        paragraph = document.add_paragraph()
        for in_single_content in content:
            font_collection = in_single_content.get('style')
            font_style = font_collection.split('-')
            if len(font_style) > 1 or font_style[0] in ['r', 'b', 'i']:
                red_data = in_single_content.get('content')
                self.add_sum_style(paragraph, red_data, font_style)
            else:
                paragraph.add_run(in_single_content['content'])
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(SPACE_PT_SIZE)

    def generate_big_second(self, document, title):
        """固定的副标题"""
        company_name = title.split('股份有限公司全体股东：')[0]
        fixed = '股份有限公司全体股东：'
        paragraph_big_second_title = document.add_paragraph()
        paragraph_big_second_title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format = paragraph_big_second_title.paragraph_format
        paragraph_format.line_spacing = Pt(SPACE_PT_SIZE)

        self.add_sum_style(paragraph_big_second_title, company_name, ['r', 'b'])
        run = paragraph_big_second_title.add_run(fixed)
        run.font.bold = True

    def generate_title(self, document, title, idx):
        """生成title"""
        index = idx - 2
        title = TITLE[index] + title
        paragraph_title = document.add_paragraph()
        run = paragraph_title.add_run(title)
        run.font.bold = True
        # paragraph_title.style = 'List Number'
        paragraph_format = paragraph_title.paragraph_format
        paragraph_format.space_before = Pt(SPACE_BEFORE_SIZE)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = Pt(SPACE_PT_SIZE)
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # Pt, Inches, RGBColor, Cm
        paragraph_format.left_indent = Cm(LEFT_INDENT_PT_SIZE)

    def generate_content(self, document, content):
        """生成正文"""
        paragraph = document.add_paragraph()
        for in_single_content in content:
            font_collection = in_single_content.get('style')
            font_style = font_collection.split('-')
            # print(len(font_style))
            # print(font_style)
            if len(font_style) > 1 or font_style[0] in ['r', 'b', 'i']:
                red_data = in_single_content.get('content')
                self.add_sum_style(paragraph, red_data, font_style)
            else:
                paragraph.add_run(in_single_content['content'])
        # 添加公共样式
        self.add_style(paragraph, document)

    def add_sum_style(self, paragraph, red_data, font_style):
        """添加 json style 字体样式"""
        run = paragraph.add_run(red_data)
        for style in font_style:
            if style == "r":
                run.font.color.rgb = RGBColor(220, 20, 60)
            elif style == "b":
                run.font.bold = True
            elif style == "i":
                run.font.italic = True

    def add_style(self, paragraph, document):
        """添加公共样式"""
        paragraph.style = document.styles['Normal']
        paragraph_format = paragraph.paragraph_format
        # paragraph_format.right_indent = Inches(1)
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.first_line_indent = Inches(FIRST_LINE_INDENT_INCH_SIZE)
        paragraph_format.line_spacing = Pt(SPACE_PT_SIZE)
        paragraph_format.space_before = Pt(SPACE_BEFORE_SIZE)
        paragraph_format.space_after = Pt(0)
        paragraph.style.font.name = u'宋体'
        paragraph.style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        font = paragraph.style.font
        font.size = Pt(CONTENT_PT_FONT_SIZE)

    def generate_word_report(self, sum_dict):
        document = Document("template.docx")
        # self.generate_header(document)
        content_list = sum_dict.get("paragraphs")
        if not content_list:
            return
        for _idx, paragraph_dict in enumerate(content_list):
            title = paragraph_dict.get("title")
            if _idx > 1:
                if title:
                    self.generate_title(document, title, _idx)
                if title == CONTENT_TABLE_LOCATION:
                    self.generate_content_table(document)
                # content 生成
                # 获取所有的段落
                paragraphs = paragraph_dict.get("paragraphs")
                # 获取单个段落
                for paragraph_id, paragraph in enumerate(paragraphs):
                    # print(paragraph)
                    self.generate_content(document, paragraph)

            elif _idx == 0:
                self.generate_big_title(document, title)
                paragraphs = paragraph_dict.get("paragraphs")
                # 获取单个段落
                for paragraph_id, paragraph in enumerate(paragraphs):
                    self.generate_head_content(document, paragraph)

            else:
                self.generate_big_second(document, title)

        self.generate_footer_table(document, sum_dict)

        return document


if __name__ == '__main__':
    logging.basicConfig(format='[%(asctime)s] - %(filename)s[line:%(lineno¬)d] - %(levelname)s: %(message)s',
                        level=logging.WARNING,
                        )
    a = GenerateWordReport()

    with open("new.json", 'r') as f:
        sum_dict = json.load(f)
        document = a.generate_word_report(sum_dict)
    document.save("demo.docx")
